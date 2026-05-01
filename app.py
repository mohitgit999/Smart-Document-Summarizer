from flask import Flask, render_template, request, jsonify, send_file
from openai import OpenAI
import os
import io
from dotenv import load_dotenv
from pypdf import PdfReader
from docx import Document
from fpdf import FPDF

# ─── Load environment variables from .env file ───────────────────────────────
load_dotenv()

app = Flask(__name__)

# ─── Initialize OpenAI client ────────────────────────────────────────────────
try:
    client = OpenAI(
        api_key=os.getenv("GROQ_API_KEY") or os.getenv("OPENAI_API_KEY"),
        base_url="https://api.groq.com/openai/v1"
    )
except Exception as e:
    print(f"Client init error: {e}")
    client = None

def extract_text_from_file(file):
    filename = file.filename.lower()
    text = ""
    try:
        if filename.endswith(".txt"):
            text = file.read().decode("utf-8")
        elif filename.endswith(".pdf"):
            reader = PdfReader(file)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        elif filename.endswith(".docx"):
            doc = Document(file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            return None # Unsupported file type
    except Exception as e:
        print(f"Error extracting text from {filename}: {e}")
        return None
    return text.strip()

# ─── Home Route ──────────────────────────────────────────────────────────────
@app.route("/")
def index():
    return render_template("index.html")

# ─── Summarize Route ─────────────────────────────────────────────────────────
@app.route("/summarize", methods=["POST"])
def summarize():
    if request.is_json:
        data = request.get_json()
        text = data.get("text", "").strip()
        style = data.get("style", "concise")
    else:
        text = request.form.get("text", "").strip()
        style = request.form.get("style", "concise")
        
        file = request.files.get("file")
        if file and file.filename:
            file_text = extract_text_from_file(file)
            if file_text is None:
                return jsonify({"error": "Failed to extract text from file. Ensure it's a valid TXT, PDF, or DOCX."}), 400
            # If there's text from file, we prepend it to any text pasted
            if file_text:
                text = (file_text + "\n\n" + text).strip()

    # ── Validation ────────────────────────────────────────────────────────────
    if not text:
        return jsonify({"error": "No text provided."}), 400

    if len(text) < 50:
        return jsonify({"error": "Text too short. Please provide at least 50 characters."}), 400

    if len(text) > 15000:
        return jsonify({"error": "Text too long. Maximum 15,000 characters allowed."}), 400

    # ── Build Prompt Based on Style ───────────────────────────────────────────
    style_prompts = {
        "concise": "Summarize the following text in 2-3 clear, concise sentences. Capture the key points only.",
        "detailed": "Provide a detailed summary of the following text in 5-7 sentences. Cover all major points and important details.",
        "bullet": "Summarize the following text as 5-7 bullet points. Each bullet should be a key insight or fact."
    }

    system_prompt = style_prompts.get(style, style_prompts["concise"])

    # ── Call OpenAI API ───────────────────────────────────────────────────────
    if not client:
        return jsonify({"error": "OpenAI/Groq client not initialized. Check API keys."}), 500

    try:
        response = client.chat.completions.create(
            model="llama-3.1-8b-instant", # Updated to a supported Groq model
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": text}
            ],
            max_tokens=500,
            temperature=0.5
        )

        summary = response.choices[0].message.content.strip()
        word_count_original = len(text.split())
        word_count_summary = len(summary.split())
        reduction = round((1 - word_count_summary / word_count_original) * 100)

        return jsonify({
            "summary": summary,
            "stats": {
                "original_words": word_count_original,
                "summary_words": word_count_summary,
                "reduction_percent": max(reduction, 0)
            }
        })

    except Exception as e:
        return jsonify({"error": f"OpenAI API Error: {str(e)}"}), 500

@app.route("/download", methods=["POST"])
def download():
    data = request.get_json()
    summary = data.get("summary", "")
    format_type = data.get("format", "txt")
    
    if not summary:
        return jsonify({"error": "No content to download"}), 400

    try:
        if format_type == "txt":
            buffer = io.BytesIO()
            buffer.write(summary.encode("utf-8"))
            buffer.seek(0)
            return send_file(buffer, as_attachment=True, download_name="summary.txt", mimetype="text/plain")

        elif format_type == "docx":
            doc = Document()
            doc.add_heading("AI Document Summary", 0)
            doc.add_paragraph(summary)
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return send_file(buffer, as_attachment=True, download_name="summary.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        elif format_type == "pdf":
            pdf = FPDF()
            pdf.add_page()
            # fpdf2 uses standard fonts by default. Unicode might need a font file.
            # We'll use a standard font for now.
            pdf.set_font("helvetica", size=12)
            pdf.multi_cell(0, 10, txt=summary)
            
            buffer = io.BytesIO()
            # fpdf2 output() returns bytes if no dest is provided
            pdf_bytes = pdf.output()
            buffer.write(pdf_bytes)
            buffer.seek(0)
            return send_file(buffer, as_attachment=True, download_name="summary.pdf", mimetype="application/pdf")

    except Exception as e:
        print(f"Download error: {e}")
        return jsonify({"error": str(e)}), 500

    return jsonify({"error": "Invalid format"}), 400


# ─── Run the App ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    app.run(debug=True, port=5000)
